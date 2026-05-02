from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Create document
doc = Document()

# TITLE PAGE
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_run = title.add_run('AI-Driven Health: A Web App for Enhanced Healthcare Queries and Nutrition Analysis')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(27, 73, 101)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_run = subtitle.add_run('COMPREHENSIVE REQUIREMENTS DOCUMENT')
subtitle_run.font.size = Pt(16)
subtitle_run.font.bold = True
subtitle_run.font.color.rgb = RGBColor(63, 142, 252)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
meta.add_run(f'Document Version: 1.0\nGenerated: {datetime.now().strftime("%B %d, %Y")}\nStatus: Production Ready').font.size = Pt(11)

doc.add_page_break()

# TABLE OF CONTENTS
doc.add_heading('TABLE OF CONTENTS', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Problem Statement',
    '3. Project Objectives',
    '4. System Architecture',
    '5. Functional Requirements',
    '6. Non-Functional Requirements',
    '7. Technical Stack',
    '8. Database Design',
    '9. Security Implementation',
    '10. Testing & Quality Assurance',
    '11. Deployment & Maintenance',
    '12. Future Enhancements'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# EXECUTIVE SUMMARY
doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
doc.add_paragraph(
    'The Medical Chatbot is an AI-powered web application that revolutionizes healthcare accessibility '
    'by providing intelligent medical guidance, nutritional analysis, and personalized health recommendations. '
    'The system leverages Google Generative AI (Gemini API) to deliver 24/7 healthcare support while maintaining '
    'enterprise-grade security and data privacy.'
)

doc.add_heading('1.1 Project Vision', level=2)
doc.add_paragraph(
    'Democratize access to health information by providing an intelligent, AI-assisted platform that enables '
    'users to receive preliminary medical guidance, track health conditions, analyze food nutrition, and receive '
    'personalized recommendations—all while maintaining the highest standards of security and privacy.'
)

doc.add_heading('1.2 Key Highlights', level=2)
highlights = [
    'Multi-language support (15+ languages)',
    'Real-time AI-powered medical responses',
    'Computer vision-based food nutrition analysis',
    'Personalized recommendations based on detected health conditions',
    'Comprehensive feedback system with admin dashboard',
    'Enterprise-grade security with bcrypt password hashing',
    'Per-user rate limiting to prevent API abuse',
    'Chat history and analysis tracking',
    'Health insights dashboard with metrics',
    'Responsive design for mobile and desktop'
]
for highlight in highlights:
    doc.add_paragraph(highlight, style='List Bullet')

doc.add_page_break()

# PROBLEM STATEMENT
doc.add_heading('2. PROBLEM STATEMENT', level=1)

doc.add_heading('2.1 Healthcare Accessibility Challenge', level=2)
doc.add_paragraph(
    'Healthcare professionals are often unavailable for quick consultations. Users lack 24/7 access to preliminary '
    'medical guidance, resulting in delayed health decisions and increased anxiety about their conditions.'
)

doc.add_heading('2.2 Health Condition Management Gap', level=2)
doc.add_paragraph(
    'Users struggle to track their health conditions across multiple conversations. There is no integrated system '
    'to detect health conditions from user queries and provide personalized medical advice.'
)

doc.add_heading('2.3 Dietary Analysis Challenges', level=2)
doc.add_paragraph(
    'Users have no easy, automated way to assess the nutritional value of their food. Current solutions require '
    'manual input or specialized knowledge, making healthy eating choices difficult.'
)

doc.add_heading('2.4 Performance Feedback Gap', level=2)
doc.add_paragraph(
    'Healthcare applications lack user feedback mechanisms. Administrators cannot monitor user satisfaction or '
    'understand user pain points for continuous improvement.'
)

doc.add_heading('2.5 Security and Privacy Concerns', level=2)
doc.add_paragraph(
    'Health data is highly sensitive. Systems require robust encryption, secure API key management, and rate '
    'limiting to protect against abuse and unauthorized access.'
)

doc.add_page_break()

# PROJECT OBJECTIVES
doc.add_heading('3. PROJECT OBJECTIVES', level=1)

doc.add_heading('3.1 Primary Objectives', level=2)
objectives = [
    'Develop an AI-powered medical chatbot using Google Generative AI',
    'Implement automatic health condition detection',
    'Create nutrition analysis system using computer vision',
    'Develop comprehensive feedback system with admin dashboard',
    'Ensure enterprise-level security implementation',
    'Provide multi-language support for global accessibility',
    'Implement personalized recommendations based on health history'
]
for i, obj in enumerate(objectives, 1):
    doc.add_paragraph(f'{i}. {obj}', style='List Bullet')

doc.add_page_break()

# SYSTEM ARCHITECTURE
doc.add_heading('4. SYSTEM ARCHITECTURE', level=1)

doc.add_heading('4.1 Architecture Overview', level=2)
doc.add_paragraph('The system follows a 3-tier MVT (Model-View-Template) architecture using Django framework.')

doc.add_heading('4.2 Technology Layers', level=2)

doc.add_heading('4.2.1 Presentation Layer', level=3)
frontend_items = [
    'Responsive HTML5 templates with Django template engine',
    'CSS3 styling with gradient effects and animations',
    'JavaScript for speech recognition and interactions',
    'Responsive grid system for mobile-first design'
]
for item in frontend_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2.2 Business Logic Layer', level=3)
business_items = [
    'Medical chatbot query processing and AI integration',
    'Nutrition analysis with image processing',
    'User authentication and session management',
    'Feedback collection and processing',
    'Health condition detection and personalization',
    'Rate limiting and request throttling'
]
for item in business_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2.3 Data Access Layer', level=3)
data_items = [
    'SQLite database with Django ORM',
    'User registration and authentication models',
    'Chat history and nutrition tracking',
    'Feedback and rating models',
    'Personalized recommendations storage',
    'Health insights and profiles'
]
for item in data_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# FUNCTIONAL REQUIREMENTS
doc.add_heading('5. FUNCTIONAL REQUIREMENTS', level=1)

doc.add_heading('5.1 User Management', level=2)
doc.add_paragraph('REQ-001: Users can register with name, email, mobile, and address information')
doc.add_paragraph('REQ-002: System validates email and mobile uniqueness')
doc.add_paragraph('REQ-003: Passwords are hashed using bcrypt (12 rounds)')
doc.add_paragraph('REQ-004: Users require admin activation before accessing features')
doc.add_paragraph('REQ-005: System supports login, logout, and session management')

doc.add_heading('5.2 Medical Chatbot Module', level=2)
doc.add_paragraph('REQ-006: Users can submit medical queries (max 500 characters)')
doc.add_paragraph('REQ-007: System sends queries to Gemini API with 30-second timeout')
doc.add_paragraph('REQ-008: AI responses are parsed into Advice and Reasoning sections')
doc.add_paragraph('REQ-009: System detects 12+ health conditions from queries')
doc.add_paragraph('REQ-010: Chat history is saved and displayed in reverse chronological order')
doc.add_paragraph('REQ-011: Users can reuse previous queries with one click')

doc.add_heading('5.3 Nutrition Bot Module', level=2)
doc.add_paragraph('REQ-012: Users can upload food images (JPEG, PNG, WebP, GIF)')
doc.add_paragraph('REQ-013: System validates image size (max 10MB)')
doc.add_paragraph('REQ-014: Images are analyzed for nutrition information')
doc.add_paragraph('REQ-015: Health percentage is displayed with visual progress bar')
doc.add_paragraph('REQ-016: Calorie information is extracted and displayed prominently')
doc.add_paragraph('REQ-017: Complete nutritional breakdown is provided')

doc.add_heading('5.4 Feedback System', level=2)
doc.add_paragraph('REQ-018: Users can submit feedback with subject, message, and rating')
doc.add_paragraph('REQ-019: Admin can view all feedbacks with statistics')
doc.add_paragraph('REQ-020: System calculates average ratings for medical bot, nutrition bot, and overall')
doc.add_paragraph('REQ-021: Feedback is timestamped and associated with user')
doc.add_paragraph('REQ-022: Admin dashboard shows total feedbacks, average rating, unique users')

doc.add_heading('5.5 Personalized Recommendations', level=2)
doc.add_paragraph('REQ-023: System generates recommendations based on detected conditions')
doc.add_paragraph('REQ-024: Recommendations are categorized (medical, nutrition, lifestyle, exercise)')
doc.add_paragraph('REQ-025: Recommendations have priority levels (high, medium, low)')
doc.add_paragraph('REQ-026: System avoids duplicate recommendations')

doc.add_heading('5.6 Health Insights Dashboard', level=2)
doc.add_paragraph('REQ-027: System calculates overall health score (0-100)')
doc.add_paragraph('REQ-028: Health status is displayed (Excellent, Good, Fair, Needs Attention)')
doc.add_paragraph('REQ-029: Activity streak is tracked')
doc.add_paragraph('REQ-030: Top health concerns are identified')

doc.add_page_break()

# NON-FUNCTIONAL REQUIREMENTS
doc.add_heading('6. NON-FUNCTIONAL REQUIREMENTS', level=1)

doc.add_heading('6.1 Security Requirements', level=2)
security_reqs = [
    'NFR-001: Passwords hashed with bcrypt (12 salt rounds)',
    'NFR-002: API keys stored in .env file (never in source code)',
    'NFR-003: CSRF tokens required for all POST requests',
    'NFR-004: All user inputs validated and sanitized',
    'NFR-005: SQL injection prevented using Django ORM',
    'NFR-006: XSS attacks prevented with template auto-escaping',
    'NFR-007: File uploads validated for type and size',
    'NFR-008: Rate limiting: 10 medical queries/min per user',
    'NFR-009: Rate limiting: 5 nutrition analyses/min per user',
    'NFR-010: API timeouts set to 30 seconds maximum'
]
for req in security_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_heading('6.2 Performance Requirements', level=2)
doc.add_paragraph('NFR-011: Medical chatbot response time < 5 seconds (90% of requests)')
doc.add_paragraph('NFR-012: Nutrition analysis response time < 8 seconds')
doc.add_paragraph('NFR-013: Page load time < 3 seconds')
doc.add_paragraph('NFR-014: Chat history loads within 2 seconds for 100+ entries')
doc.add_paragraph('NFR-015: Database queries use proper indexing')

doc.add_heading('6.3 Reliability Requirements', level=2)
doc.add_paragraph('NFR-016: System uptime 99% (excluding maintenance)')
doc.add_paragraph('NFR-017: Database backups performed daily')
doc.add_paragraph('NFR-018: Comprehensive error handling with user-friendly messages')
doc.add_paragraph('NFR-019: Logging captures all errors for debugging')

doc.add_heading('6.4 Usability Requirements', level=2)
doc.add_paragraph('NFR-020: Responsive design for mobile (320px), tablet (768px), desktop (1920px)')
doc.add_paragraph('NFR-021: Color scheme contrast minimum 4.5:1')
doc.add_paragraph('NFR-022: Clear error messages and actionable feedback')
doc.add_paragraph('NFR-023: Intuitive navigation with clear labels')

doc.add_page_break()

# TECHNICAL STACK
doc.add_heading('7. TECHNICAL STACK', level=1)

doc.add_heading('7.1 Backend', level=2)
doc.add_paragraph('Python 3.8+ - Programming language')
doc.add_paragraph('Django 5.1+ - Web framework')
doc.add_paragraph('SQLite - Development database (PostgreSQL for production)')
doc.add_paragraph('bcrypt 4.1.3 - Password hashing')
doc.add_paragraph('python-decouple 3.8 - Environment variable management')

doc.add_heading('7.2 Frontend', level=2)
doc.add_paragraph('HTML5 - Markup structure')
doc.add_paragraph('CSS3 - Styling and responsive design')
doc.add_paragraph('JavaScript (Vanilla) - Interactive features')
doc.add_paragraph('Font Awesome - UI icons')

doc.add_heading('7.3 External APIs', level=2)
doc.add_paragraph('Google Generative AI (Gemini 2.5 Flash) - Medical queries and nutrition analysis')
doc.add_paragraph('Google Cloud Speech-to-Text - Voice input')

doc.add_page_break()

# DATABASE DESIGN
doc.add_heading('8. DATABASE DESIGN', level=1)

doc.add_heading('8.1 Main Tables', level=2)
doc.add_paragraph('UserRegistrations - User accounts and authentication')
doc.add_paragraph('MedicalChatHistory - Medical queries and responses')
doc.add_paragraph('NutritionAnalysisHistory - Food analysis records')
doc.add_paragraph('UserFeedback - User feedback and ratings')
doc.add_paragraph('PersonalizedRecommendations - Health recommendations')
doc.add_paragraph('UserHealthProfile - Detected health conditions')
doc.add_paragraph('HealthInsight - User health metrics and insights')

doc.add_heading('8.2 Relationships', level=2)
doc.add_paragraph('All tables use Foreign Key to UserRegistrations (CASCADE DELETE)')
doc.add_paragraph('One-to-Many relationships for history and recommendations')
doc.add_paragraph('One-to-One relationships for health profile and insights')

doc.add_heading('8.3 Indexing Strategy', level=2)
doc.add_paragraph('user_id indexed in all tables (Foreign Key index)')
doc.add_paragraph('created_at indexed in history tables for sorting')
doc.add_paragraph('status indexed in feedback table for filtering')
doc.add_paragraph('loginid indexed in UserRegistrations for authentication')

doc.add_page_break()

# SECURITY IMPLEMENTATION
doc.add_heading('9. SECURITY IMPLEMENTATION', level=1)

doc.add_heading('9.1 Password Security', level=2)
doc.add_paragraph('• Implementation: bcrypt with 12 salt rounds')
doc.add_paragraph('• Applied to: User registration and login')
doc.add_paragraph('• Backward compatibility: Auto-migration of plain-text passwords')

doc.add_heading('9.2 API Key Management', level=2)
doc.add_paragraph('• Implementation: Environment variables via python-decouple')
doc.add_paragraph('• Storage: .env file (excluded from git)')
doc.add_paragraph('• Usage: Loaded at application startup')

doc.add_heading('9.3 Rate Limiting', level=2)
doc.add_paragraph('• Implementation: Cache-based per-user request tracking')
doc.add_paragraph('• Medical Chatbot: 10 requests per minute')
doc.add_paragraph('• Nutrition Bot: 5 requests per minute')
doc.add_paragraph('• Feedback: 5 submissions per minute')
doc.add_paragraph('• Admin Login: 5 attempts per minute')

doc.add_heading('9.4 Input Validation', level=2)
doc.add_paragraph('• Query length validation (max 500 characters)')
doc.add_paragraph('• File size validation (max 10MB)')
doc.add_paragraph('• File type validation (JPEG, PNG, WebP, GIF)')
doc.add_paragraph('• Feedback content validation')

doc.add_page_break()

# TESTING
doc.add_heading('10. TESTING & QUALITY ASSURANCE', level=1)

doc.add_heading('10.1 Unit Testing', level=2)
doc.add_paragraph('• Password hashing functions')
doc.add_paragraph('• Health condition detection')
doc.add_paragraph('• Rating calculation functions')
doc.add_paragraph('• Recommendation generation')

doc.add_heading('10.2 Integration Testing', level=2)
doc.add_paragraph('• User registration and login flow')
doc.add_paragraph('• Medical chatbot with Gemini API')
doc.add_paragraph('• Nutrition analysis with image upload')
doc.add_paragraph('• Feedback submission and retrieval')

doc.add_heading('10.3 Manual Testing', level=2)
doc.add_paragraph('• User interface on multiple browsers')
doc.add_paragraph('• Responsive design on various devices')
doc.add_paragraph('• Speech recognition functionality')
doc.add_paragraph('• Admin dashboard features')

doc.add_page_break()

# DEPLOYMENT
doc.add_heading('11. DEPLOYMENT & MAINTENANCE', level=1)

doc.add_heading('11.1 Deployment Steps', level=2)
doc.add_paragraph('1. Install dependencies: pip install -r requirements.txt')
doc.add_paragraph('2. Create .env file with API keys')
doc.add_paragraph('3. Run migrations: python manage.py migrate')
doc.add_paragraph('4. Create superuser: python manage.py createsuperuser')
doc.add_paragraph('5. Collect static files: python manage.py collectstatic')
doc.add_paragraph('6. Run development server: python manage.py runserver')

doc.add_heading('11.2 Production Considerations', level=2)
doc.add_paragraph('• Use PostgreSQL instead of SQLite')
doc.add_paragraph('• Enable HTTPS and security headers')
doc.add_paragraph('• Configure proper logging and monitoring')
doc.add_paragraph('• Set up automated backups')
doc.add_paragraph('• Use environment-specific settings')
doc.add_paragraph('• Deploy behind reverse proxy (Nginx)')

doc.add_heading('11.3 Maintenance', level=2)
doc.add_paragraph('• Regular security updates and patches')
doc.add_paragraph('• Database optimization and cleanup')
doc.add_paragraph('• Monitor API usage and quota')
doc.add_paragraph('• Review and analyze user feedback')

doc.add_page_break()

# FUTURE ENHANCEMENTS
doc.add_heading('12. FUTURE ENHANCEMENTS', level=1)
doc.add_paragraph('• Mobile app development (iOS/Android)')
doc.add_paragraph('• Advanced health analytics and reporting')
doc.add_paragraph('• Integration with medical devices')
doc.add_paragraph('• Telemedicine consultation features')
doc.add_paragraph('• Medication interaction checker')
doc.add_paragraph('• Appointment scheduling with doctors')
doc.add_paragraph('• Prescription management')
doc.add_paragraph('• Health records export and sharing')
doc.add_paragraph('• Machine learning-based personalization')
doc.add_paragraph('• Video consultation capability')

# Save document
doc.save('REQUIREMENTS_DOCUMENT.docx')
print("✅ Requirements document created successfully: REQUIREMENTS_DOCUMENT.docx")
